"""월간 리포트 데이터 → 편집 가능한 Word(.docx) 바이트.

양식: 제목/부제/메타 → 한 문장 요약 → 01 한 장 요약표 → 02 좋아진·아쉬운 →
03 매체별 비교표 + 해석 → 04 키워드표 + 코멘트 → 05 다음달 전략(초안).
전략은 편집용 초안 — 사람이 매월 다듬어 쓰는 것을 전제로 한다.
"""
from __future__ import annotations

import io

BRAND = (0x1E, 0x40, 0xAF)
INK = (0x1F, 0x29, 0x37)
MUT = (0x6B, 0x72, 0x80)
POS = (0x2E, 0x7A, 0x57)
NEG = (0xC2, 0x4A, 0x38)


def _next_month_label(target: str) -> str:
    y, m = (int(x) for x in target.split("-"))
    return "1월" if m == 12 else f"{m+1}월"


def build_report_docx(data: dict) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.table import WD_TABLE_ALIGNMENT

    def C(t):
        return RGBColor(*t)

    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = "맑은 고딕"
    st.font.size = Pt(10.5)
    st.font.color.rgb = C(INK)
    try:
        from docx.oxml.ns import qn
        st.element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    except Exception:
        pass

    def run(p, text, bold=False, color=None, size=None):
        r = p.add_run(text)
        r.bold = bold
        if color:
            r.font.color.rgb = C(color)
        if size:
            r.font.size = Pt(size)
        return r

    def sec(n, title):
        p = doc.add_paragraph()
        run(p, f"{n}  ", bold=True, color=BRAND, size=13)
        run(p, title, bold=True, size=13)

    def table(headers, rows):
        t = doc.add_table(rows=1, cols=len(headers))
        t.style = "Light Grid Accent 1"
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, h in enumerate(headers):
            run(t.rows[0].cells[i].paragraphs[0], h, bold=True)
        for r0 in rows:
            cells = t.add_row().cells
            for i, x in enumerate(r0):
                run(cells[i].paragraphs[0], x, bold=(i == 0))

    tl = data["target_label"]
    cl = data["compare_label"]

    # 제목
    p = doc.add_paragraph()
    run(p, "keek 온라인 광고 월간 리포트", bold=True, size=20)
    p = doc.add_paragraph()
    run(p, f"{cl} → {tl}  ·  성과 비교 분석", bold=True, color=BRAND, size=12.5)
    p = doc.add_paragraph()
    run(p, "대상  Meta · 네이버 파워링크 · 네이버 플레이스     |     지표 기준  매출·ROAS·CTR·CPC     |     데이터  광고 API + 자사몰 실측",
        color=MUT, size=9)
    p = doc.add_paragraph()
    run(p, "한 문장 요약 — ", bold=True, color=BRAND)
    run(p, data["headline"])

    # 01
    sec("01", "한 장 요약 · 계정 전체")
    table(["지표", cl, tl, "증감"],
          [[k["label"], k["a"], k["b"], k["delta"]["text"]] for k in data["kpi_rows"]])

    # 02
    sec("02", "좋아진 것 · 아쉬운 것")
    p = doc.add_paragraph()
    run(p, "◆ 개선된 지표", bold=True, color=POS, size=11.5)
    for s in data["improved"]:
        pp = doc.add_paragraph(style="List Bullet")
        run(pp, s)
    p = doc.add_paragraph()
    run(p, "◆ 아쉬운 · 점검할 지표", bold=True, color=NEG, size=11.5)
    for s in data["declined"]:
        pp = doc.add_paragraph(style="List Bullet")
        run(pp, s)

    # 03
    sec("03", f"매체별 비교 ({cl} → {tl})")
    crows = []
    for c in data["channel_rows"]:
        crows.append([c["channel"],
                      f'{c["ad_cost"]} ({c["ad_cost_d"]["text"]})',
                      f'{c["impressions"]} ({c["impressions_d"]["text"]})',
                      f'{c["clicks"]} ({c["clicks_d"]["text"]})',
                      f'{c["ctr"]} ({c["ctr_d"]["text"]})',
                      f'{c["cpc"]} ({c["cpc_d"]["text"]})',
                      f'{c["ad_sales"]}',
                      f'{c["roas"]} ({c["roas_d"]["text"]})'])
    table(["채널", "광고비", "노출", "클릭", "CTR", "CPC", "광고매출", "ROAS"], crows)
    p = doc.add_paragraph()
    run(p, "해석 — ", bold=True)
    run(p, data["interpretation"])

    # 04
    if data["top_keywords"]:
        sec("04", f"네이버 검색 · 키워드 흐름 ({tl})")
        krows = [[k["keyword"], k["ad_cost"], k["clicks"], k["ctr"], k["cpc"]] for k in data["top_keywords"]]
        table(["키워드", "광고비", "클릭", "CTR", "CPC"], krows)
        if data["keyword_note"]:
            p = doc.add_paragraph()
            run(p, data["keyword_note"])

    # 05
    sec("05", f"핵심 인사이트 & {_next_month_label(data['target'])} 전략 (초안 · 편집용)")
    for i, s in enumerate(data["strategy_draft"], 1):
        p = doc.add_paragraph()
        run(p, f"{i}. ", bold=True, color=BRAND, size=11.5)
        run(p, s["title"], bold=True, size=11.5)
        pb = doc.add_paragraph()
        pb.paragraph_format.left_indent = Cm(0.6)
        run(pb, s["body"], color=MUT)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
